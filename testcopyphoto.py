'''

.JPG to .docx

@author: Benjamin Luk


##### READ ME #######

Setup:

1. Please make sure your project folder does not contain any visible .JPG file or they will be deleted !

2. One main folder containing all subfolders of .jpg in working/project/global environment

i.e. if [____] is a folder
[Trees (Main folder)] --> [Branch 1], [Branch 2], .....
                              |
                              V
               Leaf1.JPG, Leaf2.JPG, Leaf3.JPG ....


3. Create empty .docx file in global environment


'''



from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
import os
import shutil
import os.path
import math


# Inputs

projpathinput = '/Users/benjaminluk/PycharmProjects/photodocx'
treepathinput = '/Users/benjaminluk/PycharmProjects/photodocx/Trees'
widthinput = Inches(2.0)
heightinput = Inches(2.0)
documentname = 'Testphoto.docx'



'''
with open(csvname, 'r') as csv_file:
    treefolderlist = []
    csv_reader = csv.DictReader(csv_file, delimiter = ',')
    for lines in csv_reader:
        x = lines['TREE_MARK']
        treefolderlist.append(x)
    print(treefolderlist)
'''

treefolderlist = os.listdir(treepathinput)
document = Document()

# Count number of .jpg in one subfolder
def image_count(treefolder):
    y = len([name for name in os.listdir(f"{treepathinput}/{treefolder}") if name.endswith('.JPG') or name.endswith(('.jpg'))])
    return y

# copy all .JPG in one subfolder to global env
def rearrange(treefolder):

    # --------------------------------------------------------
    reorg_dir = f"{treepathinput}/{treefolder}"
    target_dir = f"{projpathinput}"
    # ---------------------------------------------------------
    for root, dirs, files in os.walk(reorg_dir):
        for name in files:
            subject = root + "/" + name
            n = 1;
            name_orig = name
            while os.path.exists(target_dir + "/" + name):
                name = "duplicate_" + str(n) + "_" + name_orig;
                n = n + 1
            newfile = target_dir + "/" + name;
            shutil.copy(subject, newfile)



# removes all .JPGs in global environment
def remove():
    dir_name = f"{projpathinput}/"
    test = os.listdir(dir_name)

    for item in test:
        if item.endswith(".JPG") or item.endswith('.jpg'):
            os.remove(os.path.join(dir_name, item))
        if item.endswith('.DS_Store'):
            os.remove(os.path.join(dir_name, item))

# function to rename each image to number ordered
def rename(treefolder):
    x = list(range(0, image_count(treefolder)))
    i = 0
    for item in sorted(os.listdir(f"{projpathinput}")):
        if item.endswith('.JPG'):
            os.rename(item, f'{x[i]}.JPG')
            i += 1

# function to change current page to landscape
def change_orientation():
    current_section = document.sections[0]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.CONTINUOUS)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section


# Insert one subfolder of .jpg into a page in .docx
def table_picture_insert(treefolder = 'Tree_01', n_col = 4, width = widthinput, height = heightinput):

    # Rotate current page to landscape
    change_orientation()

    # add images from one folder to global environment and rename in order
    rearrange(treefolder)
    rename(treefolder)

    count = image_count(treefolder)
    image_range = range(0, count)

    templist = []

    # create list of images in '0.JPG, 1.JPG,...' for each folder
    for i in image_range:
        image_string = f'{i}.JPG'
        templist.append(image_string)

    # create table with corresponding rows and columns
    tbl = document.add_table(rows = 0, cols = n_col)

    # number of rows required
    n_row = math.ceil(len(templist))

    index = 0

    # add row to table by iterating through the number of row required
    for row in range(0, n_row):
        row_cells = tbl.add_row()
        # add image by iterating through each column in a row
        for column in range(0, n_col):
            paragraph = row_cells.cells[column].paragraphs[0]
            run = paragraph.add_run()
            if index < len(templist):
                run.add_picture(templist[index], width = widthinput, height = heightinput)
                index += 1

    # add new page
    document.add_section()

    # remove .JPG in global environment
    remove()


# Repeat .jpg insertion for each subfolder into .docx
def auto_collate(treefolderlist):
    treelist = os.listdir(f'{treepathinput}')
    for i in treefolderlist:
        if i != '.DS_Store' and i in treelist:
            table_picture_insert(treefolder= i)



# Run process

auto_collate(treefolderlist)

remove()

document.save(documentname)